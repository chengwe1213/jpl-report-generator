#!/usr/bin/env python3
"""
Fill JPL Report Template from Excel Data

Reads data from 'To Word 1' and 'To Word 2' tabs in Excel and populates the Word template.
- 'To Word 1': Sample information table (nested in main table)
- 'To Word 2': FlowCam quantification table (Table 1 in document)

If more than 6 samples exist in 'To Word 1', additional tables are generated.
For 'To Word 2', rows are added dynamically to accommodate all samples.
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import sys
import os
import re
import glob
import math
import zipfile
import shutil
import tempfile

# Optional imports for image reordering
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


def read_excel_data_transposed(excel_path: str, sheet_name: str = "To Word 1") -> pd.DataFrame:
    """Read and transpose the Excel data to get samples as rows (for 'To Word 1' format)."""
    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
    
    # First column contains field names, remaining columns are samples
    df = df.set_index(0).T
    df = df.reset_index(drop=True)
    
    # Clean column names
    df.columns = [str(col).strip() if pd.notna(col) else '' for col in df.columns]
    
    return df


def read_excel_data_rows(excel_path: str, sheet_name: str = "To Word 2") -> pd.DataFrame:
    """Read Excel data where first row is header and subsequent rows are samples (for 'To Word 2' format)."""
    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=0)
    
    # Clean column names
    df.columns = [str(col).strip() if pd.notna(col) else '' for col in df.columns]
    
    return df


def get_sample_table(doc: Document):
    """Find the nested sample table in the document."""
    # The sample table is nested in Table 0, Row 8, Cell 0
    main_table = doc.tables[0]
    cell = main_table.rows[8].cells[0]
    
    if cell.tables:
        return cell.tables[0]
    return None


def get_sample_table_parent_cell(doc: Document):
    """Get the parent cell containing the sample table."""
    main_table = doc.tables[0]
    return main_table.rows[8].cells[0]


# Mapping from Excel column names to Word table row indices
# Note: Row 6 has merged cells for "Orientation" label, Row 7 has the data cells
FIELD_TO_ROW = {
    "JPL Sample #": 0,
    "Lot No / Formulation No": 1,
    "Container": 2,
    "Fill Volume": 3,
    "Storage time": 4,
    "Storage temp": 5,
    "Orientation": 7,  # Data is in row 7, not row 6 (row 6 is the label with merged cells)
    "Lims No": 8,  # Maps to sLIMS No
    "Remark": 9,
}

# Additional fields that might be in Excel but not in the main sample table
EXTRA_FIELDS = ["EP result", "Seidenader"]


def fill_ep_seidenader_row(doc_path: str, ep_values: list, seidenader_values: list):
    """
    Fill the EP/Seidenader row in the Sample Information Table.
    
    Selects the appropriate comboBox option based on data availability:
    - Both EP and Seidenader have values -> "EP*/Seidenader result", fill "EP/Seidenader"
    - Only EP has value -> "EP result*", fill EP value
    - Only Seidenader has value -> "Seidenader result", fill Seidenader value
    - Both NA/empty -> Leave as "Choose an item."
    
    Args:
        doc_path: Path to the Word document
        ep_values: List of EP result values for each sample
        seidenader_values: List of Seidenader values for each sample
    """
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extract the docx
        extract_dir = os.path.join(temp_dir, 'docx_extracted')
        with zipfile.ZipFile(doc_path, 'r') as z:
            z.extractall(extract_dir)
        
        # Read document.xml
        doc_xml_path = os.path.join(extract_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Determine which option to select based on data availability
        def has_value(val):
            """Check if value is not NA/empty."""
            if val is None:
                return False
            if pd.isna(val):
                return False
            val_str = str(val).strip().lower()
            return val_str not in ('', 'na', 'n/a', 'nan', 'none')
        
        # Check all samples to determine the pattern
        has_ep = any(has_value(v) for v in ep_values)
        has_seid = any(has_value(v) for v in seidenader_values)
        
        # Determine which option to select
        if has_ep and has_seid:
            selected_option = "EP*/Seidenader result"
        elif has_ep:
            selected_option = "EP result*"
        elif has_seid:
            selected_option = "Seidenader result"
        else:
            selected_option = None  # Leave as default
        
        if selected_option:
            # Replace the comboBox SDT with plain text showing the selected option
            # Pattern for EP/Seidenader comboBox
            ep_combobox_pattern = (
                r'<w:comboBox>'
                r'<w:listItem w:value="Choose an item."/>'
                r'<w:listItem w:displayText="EP result\*" w:value="EP result\*"/>'
                r'<w:listItem w:displayText="Seidenader result" w:value="Seidenader result"/>'
                r'<w:listItem w:displayText="EP\*/Seidenader result" w:value="EP\*/Seidenader result"/>'
                r'</w:comboBox>'
            )
            
            # Find the full SDT element containing this comboBox
            sdt_pattern = re.compile(
                r'<w:sdt><w:sdtPr>(?:(?!</w:sdtPr>).)*?' + ep_combobox_pattern + r'</w:sdtPr>'
                r'<w:sdtEndPr/>'
                r'<w:sdtContent>((?:(?!</w:sdtContent>).)*?)</w:sdtContent></w:sdt>',
                re.DOTALL
            )
            
            def replace_ep_sdt(match):
                sdt_content = match.group(1)
                # Replace placeholder text with selected option
                modified = re.sub(
                    r'<w:t[^>]*>[^<]*</w:t>',
                    f'<w:t>{selected_option}</w:t>',
                    sdt_content,
                    count=1
                )
                # Remove PlaceholderText style
                modified = re.sub(r'<w:rStyle w:val="PlaceholderText"/>', '', modified)
                return modified
            
            content = sdt_pattern.sub(replace_ep_sdt, content)
            print(f"  Selected EP/Seidenader option: {selected_option}")
        
        # Now fill the EP/Seidenader values in the sample columns (cells 2-7 of row 8)
        # Find the table containing the EP/Seidenader row
        # The row is identified by having the EP comboBox or the selected text
        
        # Build the values to fill based on the selected option
        fill_values = []
        for i in range(len(ep_values)):
            ep_val = ep_values[i] if i < len(ep_values) else None
            seid_val = seidenader_values[i] if i < len(seidenader_values) else None
            
            if selected_option == "EP*/Seidenader result":
                # Format: "EP_value/Seidenader_value"
                ep_str = str(ep_val) if has_value(ep_val) else "NA"
                seid_str = str(seid_val) if has_value(seid_val) else "NA"
                fill_values.append(f"{ep_str}/{seid_str}")
            elif selected_option == "EP result*":
                fill_values.append(str(ep_val) if has_value(ep_val) else "")
            elif selected_option == "Seidenader result":
                fill_values.append(str(seid_val) if has_value(seid_val) else "")
            else:
                fill_values.append("")
        
        # Write modified document.xml
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Repackage the docx
        with zipfile.ZipFile(doc_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, extract_dir)
                    z.write(file_path, arc_name)
        
        # Now use python-docx to fill the values with proper formatting
        if fill_values and any(v for v in fill_values):
            doc = Document(doc_path)
            # Find all sample tables (nested in main table, row 8, cell 0)
            main_table = doc.tables[0]
            parent_cell = main_table.rows[8].cells[0]
            
            total_filled = 0
            for table_idx, sample_table in enumerate(parent_cell.tables):
                ep_row_idx = 7  # 0-indexed, row 8 in 1-indexed
                start_sample = table_idx * 6
                
                if len(sample_table.rows) > ep_row_idx:
                    row = sample_table.rows[ep_row_idx]
                    
                    # Format the header cell (column 0) - Roche Sans, 10pt, Bold
                    header_cell = row.cells[0]
                    header_cell.text = ""
                    para = header_cell.paragraphs[0]
                    run = para.add_run(selected_option)
                    run.font.name = "Roche Sans"
                    run.font.size = Pt(10)
                    run.bold = True
                    
                    # Fill value cells with proper formatting - Roche Sans, 10pt, not bold
                    for col_idx in range(6):
                        value_idx = start_sample + col_idx
                        if value_idx >= len(fill_values):
                            break
                        cell_idx = col_idx + 1
                        if cell_idx < len(row.cells):
                            value_cell = row.cells[cell_idx]
                            value_cell.text = ""
                            para = value_cell.paragraphs[0]
                            run = para.add_run(str(fill_values[value_idx]))
                            run.font.name = "Roche Sans"
                            run.font.size = Pt(10)
                            run.bold = False
                            total_filled += 1
            
            doc.save(doc_path)
            print(f"  Filled EP/Seidenader values for {total_filled} samples across {len(parent_cell.tables)} tables")
        
        return selected_option, fill_values
        
    except Exception as e:
        print(f"  ERROR filling EP/Seidenader row: {e}")
        import traceback
        traceback.print_exc()
        return None, []
        
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def update_flowcam_visual_inspection_header(doc_path: str, header_text: str, values: list):
    """
    Update the FlowCam Table 'Visual Inspection (EP)' header and fill column values.
    Preserves the original formatting (bold, font).
    
    Args:
        doc_path: Path to the Word document
        header_text: New header text (e.g., "EP*/Seidenader result")
        values: List of values to fill in the Visual Inspection column
    """
    doc = Document(doc_path)
    flowcam_table = doc.tables[5]  # FlowCam table is table index 5
    
    # Update header in rows 0, 1, 2 (all three header rows have the same text in column 6)
    for row_idx in range(3):
        cell = flowcam_table.rows[row_idx].cells[6]
        # Preserve formatting by updating the first run's text
        if cell.paragraphs and cell.paragraphs[0].runs:
            # Keep the first run's formatting, update text
            cell.paragraphs[0].runs[0].text = header_text
            # Remove any additional runs
            for run in cell.paragraphs[0].runs[1:]:
                run.text = ""
        else:
            cell.text = header_text
    
    # Fill values in data rows (starting from row 3)
    for i, value in enumerate(values):
        row_idx = i + 3  # Skip 3 header rows
        if row_idx < len(flowcam_table.rows):
            flowcam_table.rows[row_idx].cells[6].text = str(value) if value else ""
    
    doc.save(doc_path)
    print(f"  Updated FlowCam header to: {header_text}")
    print(f"  Filled Visual Inspection values for {len(values)} samples")


def remove_dropdown_from_orientation_row(doc_path: str, orientation_values: list):
    """
    Remove dropdown controls from Orientation row and fill with values.
    
    This modifies the document XML directly because the SDT elements
    are not accessible through python-docx.
    
    Args:
        doc_path: Path to the Word document
        orientation_values: List of orientation values for each sample
    """
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extract the docx
        extract_dir = os.path.join(temp_dir, 'docx_extracted')
        with zipfile.ZipFile(doc_path, 'r') as z:
            z.extractall(extract_dir)
        
        # Read document.xml
        doc_xml_path = os.path.join(extract_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Pattern for Orientation dropdown options
        orientation_dropdown = (
            r'<w:dropDownList>'
            r'<w:listItem w:value="Choose an item."/>'
            r'<w:listItem w:displayText="n/a" w:value="n/a"/>'
            r'<w:listItem w:displayText="Upright" w:value="Upright"/>'
            r'<w:listItem w:displayText="Inverted" w:value="Inverted"/>'
            r'<w:listItem w:displayText="Horizontal" w:value="Horizontal"/>'
            r'</w:dropDownList>'
        )
        
        # The SDT structure varies between templates
        # Some have <w:sdtEndPr/>, some don't
        # Pattern: <w:sdt>...<w:dropDownList>...</w:dropDownList>...</w:sdtPr>[optional: <w:sdtEndPr/>]<w:sdtContent>...</w:sdtContent></w:sdt>
        sdt_pattern = re.compile(
            r'<w:sdt><w:sdtPr>(?:(?!</w:sdtPr>).)*?' + orientation_dropdown + r'</w:sdtPr>'
            r'(?:<w:sdtEndPr/>)?'
            r'<w:sdtContent>(.*?)</w:sdtContent></w:sdt>',
            re.DOTALL
        )
        
        value_idx = 0
        def replace_sdt(match):
            nonlocal value_idx
            sdt_content = match.group(1)
            
            # Get the value to insert
            if value_idx < len(orientation_values):
                value = str(orientation_values[value_idx])
            else:
                adjusted_idx = value_idx % len(orientation_values) if orientation_values else 0
                value = str(orientation_values[adjusted_idx]) if orientation_values else ""
            value_idx += 1
            
            # Replace the placeholder text with actual value
            modified = re.sub(
                r'<w:t[^>]*>[^<]*</w:t>',
                f'<w:t>{value}</w:t>',
                sdt_content,
                count=1
            )
            
            # Remove PlaceholderText style
            modified = re.sub(r'<w:rStyle w:val="PlaceholderText"/>', '', modified)
            
            return modified
        
        # Replace all Orientation dropdowns
        content = sdt_pattern.sub(replace_sdt, content)
        
        # Write modified document.xml
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Repackage the docx
        with zipfile.ZipFile(doc_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, extract_dir)
                    z.write(file_path, arc_name)
        
        print(f"  Removed {value_idx} dropdown controls from Orientation fields")
        return True
        
    except Exception as e:
        print(f"  ERROR removing dropdowns: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def fill_sample_table(table, samples_df: pd.DataFrame, start_sample: int = 0):
    """
    Fill a sample table with data from the DataFrame.
    
    Args:
        table: The Word table to fill
        samples_df: DataFrame with sample data
        start_sample: Starting index in the DataFrame
    """
    num_samples = min(6, len(samples_df) - start_sample)
    
    for sample_idx in range(num_samples):
        df_idx = start_sample + sample_idx
        if df_idx >= len(samples_df):
            break
            
        sample = samples_df.iloc[df_idx]
        col_idx = sample_idx + 1  # Column 0 is the label column
        
        for field_name, row_idx in FIELD_TO_ROW.items():
            # Find matching column in DataFrame (handle variations)
            value = None
            for col in samples_df.columns:
                if field_name.lower() in col.lower() or col.lower() in field_name.lower():
                    value = sample.get(col)
                    break
            
            # Special handling for Storage temp (might have different naming)
            if field_name == "Storage temp" and value is None:
                for col in samples_df.columns:
                    if "storage" in col.lower() and "temp" in col.lower():
                        value = sample.get(col)
                        break
            
            # Special handling for Lims No / sLIMS No
            if field_name == "Lims No" and value is None:
                for col in samples_df.columns:
                    if "lims" in col.lower():
                        value = sample.get(col)
                        break
            
            if value is not None and pd.notna(value):
                try:
                    cell = table.rows[row_idx].cells[col_idx]
                    # Skip Orientation field - it will be handled separately via XML
                    if field_name == "Orientation":
                        continue
                    # Clear existing content and set new value
                    cell.text = str(value)
                except IndexError:
                    pass  # Skip if cell doesn't exist


def copy_table(table):
    """Create a deep copy of a table's XML."""
    return deepcopy(table._tbl)


def copy_row(row):
    """Create a deep copy of a table row's XML."""
    return deepcopy(row._tr)


def get_flowcam_table(doc: Document):
    """Get the FlowCam quantification table (Table 5 in document structure, 'Table 1' in content)."""
    return doc.tables[5]


# Mapping from Excel 'To Word 2' columns to FlowCam table columns
# Note: EP result (col 6) and comment (col 7) are left blank
FLOWCAM_COLUMN_MAP = {
    "JPL Sample #": 0,
    "Lot No / Formulation No": 1,  # Maps to "Batch" column
    "Storage time": 2,
    "Storage temp": 3,
    "Container": 4,
    "Orientation": 5,
}

# Columns for Particle concentration and Particle count
PARTICLE_CONCENTRATION_COLS = [8, 9, 10, 11, 12]  # ≥3, ≥5, ≥10, ≥25, ≥50
PARTICLE_COUNT_COLS = [13, 14]  # ≥100, ≥150

# Mapping from CSV filter names to table columns
# Particle concentration uses P/ML values (index 1 in CSV row after split)
CONCENTRATION_FILTERS = [
    ("ECD_bigger_3um", 8),    # ≥3
    ("ECD_bigger_5um", 9),    # ≥5
    ("ECD_bigger_10um", 10),  # ≥10
    ("ECD_bigger_25um", 11),  # ≥25
    ("ECD_bigger_50um", 12),  # ≥50
]

# Particle count uses Count values (index 0 in CSV row after split)
COUNT_FILTERS = [
    ("Length_bigger_100um", 13),  # ≥100
    ("Lenght_bigger_150um", 14),  # ≥150 (note typo in source)
]


def extract_sample_number_from_folder(folder_name: str) -> int:
    """
    Extract sample number from folder name.
    E.g., 'JPL25-0180_Trontinemab_F1_0_5_others_Horizontal_CWe_30Dec2025_processed' -> 1
    """
    match = re.search(r'_F(\d+)_', folder_name)
    if match:
        return int(match.group(1))
    return None


def extract_batch_from_folder(folder_name: str) -> str:
    """
    Extract batch/lot name from folder name.
    E.g., 'JPL25-0172_Abx_BS2003ES01_ 2W_5_Vial_Inverted_CWe_29Dec2025_processed' -> 'BS2003ES01'
    E.g., 'JPL25-0172_Abx_EXP-25-AC3979-EP_ 2W_5_Vial_Inverted_CWe_29Dec2025_processed' -> 'EXP-25-AC3979-EP'
    """
    # Split by underscore and find the batch part (usually 3rd element)
    parts = folder_name.split('_')
    if len(parts) >= 3:
        # The batch is typically after the product name
        # Handle cases like "EXP-25-AC3979-EP" which contains hyphens
        batch = parts[2]
        return batch
    return None


def parse_csv_for_particle_data(csv_path: str) -> dict:
    """
    Parse a summary_export.csv file and extract particle concentration and count data.
    
    Returns dict with:
        - concentration: {3: value, 5: value, 10: value, 25: value, 50: value}
        - count: {100: value, 150: value}
    """
    result = {
        "concentration": {},  # P/ML values for ≥3, ≥5, ≥10, ≥25, ≥50
        "count": {},          # Count values for ≥100, ≥150
    }
    
    try:
        with open(csv_path, 'r') as f:
            lines = f.readlines()
        
        for line in lines:
            parts = line.strip().split(',')
            if len(parts) < 3:
                continue
            
            filter_name = parts[0].strip()
            
            # Check concentration filters (use P/ML value - index 2)
            for csv_filter, col_idx in CONCENTRATION_FILTERS:
                if filter_name == csv_filter:
                    try:
                        value = float(parts[2])
                        # Round up to integer
                        result["concentration"][col_idx] = math.ceil(value)
                    except (ValueError, IndexError):
                        result["concentration"][col_idx] = 0
            
            # Check count filters (use Count value - index 1)
            for csv_filter, col_idx in COUNT_FILTERS:
                if filter_name == csv_filter:
                    try:
                        value = float(parts[1])
                        # Round up to integer
                        result["count"][col_idx] = math.ceil(value)
                    except (ValueError, IndexError):
                        result["count"][col_idx] = 0
    
    except Exception as e:
        print(f"  WARNING: Error reading {csv_path}: {e}")
    
    return result


def find_particle_data_folders(base_path: str, sample_order: list = None) -> dict:
    """
    Find all processed folders and extract particle data.
    
    Args:
        base_path: Path to search for processed folders
        sample_order: Optional list of batch/lot names in sample order (from Excel)
                     If provided, maps batch names to sample numbers
    
    Returns dict mapping sample number to particle data.
    """
    sample_data = {}
    
    # Find all folders matching the pattern
    pattern = os.path.join(base_path, "*_processed")
    folders = glob.glob(pattern)
    
    if not folders:
        # Try with space in pattern
        pattern = os.path.join(base_path, "*processed*")
        folders = [f for f in glob.glob(pattern) if os.path.isdir(f)]
    
    # Create batch-to-sample mapping if sample_order provided
    batch_to_sample = {}
    if sample_order:
        for i, batch in enumerate(sample_order, 1):
            # Normalize batch name for matching (handle variations like "LTA7TA1003/BS25080738")
            batch_key = batch.split('/')[0] if '/' in batch else batch
            batch_to_sample[batch_key.upper()] = i
    
    for folder in sorted(folders):
        folder_name = os.path.basename(folder)
        
        # Try to extract sample number using F# pattern first
        sample_num = extract_sample_number_from_folder(folder_name)
        
        # If no F# pattern, try batch name matching
        if sample_num is None and sample_order:
            batch = extract_batch_from_folder(folder_name)
            if batch:
                # Try to match batch to sample order
                batch_upper = batch.upper()
                sample_num = batch_to_sample.get(batch_upper)
                
                # Also try partial matching for complex batch names
                if sample_num is None:
                    for batch_key, num in batch_to_sample.items():
                        if batch_upper in batch_key or batch_key in batch_upper:
                            sample_num = num
                            break
        
        if sample_num is None:
            continue
        
        # Find CSV file in folder
        csv_files = glob.glob(os.path.join(folder, "*.csv"))
        if not csv_files:
            csv_files = glob.glob(os.path.join(folder, "*.CSV"))
        
        if csv_files:
            csv_path = csv_files[0]  # Use first CSV found
            batch_info = extract_batch_from_folder(folder_name) or f"F{sample_num}"
            print(f"  Sample {sample_num} ({batch_info}): Reading {os.path.basename(csv_path)}")
            particle_data = parse_csv_for_particle_data(csv_path)
            sample_data[sample_num] = particle_data
    
    return sample_data


def fill_particle_data_in_table(table, particle_data: dict):
    """
    Fill particle concentration and count data in the FlowCam table.
    
    Args:
        table: The FlowCam Word table
        particle_data: Dict mapping sample number to particle data
    """
    # Data rows start at row 3 (after 3 header rows)
    for sample_num, data in particle_data.items():
        row_idx = sample_num + 2  # Sample 1 -> Row 3, Sample 2 -> Row 4, etc.
        
        if row_idx >= len(table.rows):
            print(f"  WARNING: Row {row_idx} not found for sample {sample_num}")
            continue
        
        # Fill concentration columns
        for col_idx, value in data.get("concentration", {}).items():
            try:
                table.rows[row_idx].cells[col_idx].text = str(value)
            except IndexError:
                pass
        
        # Fill count columns
        for col_idx, value in data.get("count", {}).items():
            try:
                table.rows[row_idx].cells[col_idx].text = str(value)
            except IndexError:
                pass


# ============== IMAGE REORDERING FUNCTIONS ==============

def extract_batch_from_image(image_path: str) -> str:
    """
    Use OCR to extract batch name from morphology image.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        Batch name (e.g., "F1", "BS2403SA03", etc.) or None if not found
    """
    if not HAS_OCR:
        return None
    
    try:
        img = Image.open(image_path)
        
        # Only read the top portion where the text header is
        width, height = img.size
        top_portion = img.crop((0, 0, width, min(120, height)))
        
        text = pytesseract.image_to_string(top_portion)
        
        # Try to extract batch name using regex
        # Pattern 1: "Batch : F1" or "Batch: F2"
        match = re.search(r'Batch\s*:\s*F(\d+)', text, re.IGNORECASE)
        if match:
            return f"F{match.group(1)}"
        
        # Pattern 2: "Batch : BS2403SA03" or other alphanumeric batch names
        match = re.search(r'Batch\s*:\s*([A-Za-z0-9\-]+)', text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
            
    except Exception as e:
        print(f"  WARNING: Could not read image {image_path}: {e}")
    
    return None


def identify_morphology_images(docx_path: str, temp_dir: str) -> dict:
    """
    Extract images from docx and identify which batch each belongs to.
    
    Args:
        docx_path: Path to the Word document
        temp_dir: Temporary directory to extract images
        
    Returns:
        Dict mapping image filename to batch name (string)
    """
    image_batches = {}
    
    # Extract images from docx
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if 'word/media/' in name and name.endswith('.png'):
                z.extract(name, temp_dir)
                image_path = os.path.join(temp_dir, name)
                
                batch_name = extract_batch_from_image(image_path)
                if batch_name is not None:
                    image_name = os.path.basename(name)
                    image_batches[image_name] = batch_name
                    print(f"    {image_name}: Batch {batch_name}")
    
    return image_batches


def get_image_groups_from_document(docx_path: str) -> list:
    """
    Analyze document structure to find morphology image sections.
    
    Returns list of dicts with section info and image references.
    """
    doc = Document(docx_path)
    
    sections = []
    current_section = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        
        # Detect morphology sections
        if 'images from particles' in text and 'ecd' in text and '<100' in text:
            current_section = {'type': 'small', 'name': 'ECD <100µm', 'para_idx': i, 'images': []}
            sections.append(current_section)
        elif 'images from particles' in text and '100' in text and 'length' in text:
            current_section = {'type': 'large', 'name': '≥100µm Length', 'para_idx': i, 'images': []}
            sections.append(current_section)
    
    return sections


def reorder_morphology_images(docx_path: str, output_path: str, sample_order: list = None, template_path: str = None) -> bool:
    """
    Reorder morphology images in the document according to sample sequence.
    
    This function:
    1. Extracts images from the docx
    2. Uses OCR to identify batch names
    3. Swaps the entire drawing elements (including extents) in document.xml
    4. Preserves original image sizes and quality
    
    Args:
        docx_path: Path to input Word document
        output_path: Path for output Word document
        sample_order: Optional list of batch names in desired order (from Excel)
        
    Returns:
        True if successful, False otherwise
    """
    if not HAS_OCR:
        print("WARNING: OCR libraries not available. Cannot reorder images.")
        print("  Install with: pip install pytesseract pillow")
        return False
    
    print("\n--- Reordering morphology images ---")
    
    # Create temp directory
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Step 1: Extract and identify images
        print("  Identifying images by batch name...")
        image_batches = identify_morphology_images(docx_path, temp_dir)
        
        if not image_batches:
            print("  No morphology images found with batch information")
            return False
        
        # Step 2: Create batch order mapping
        # If sample_order provided, use it; otherwise sort alphabetically
        
        def normalize_batch_name(name):
            """Normalize batch name to handle OCR errors (O vs 0, etc.)."""
            if name is None:
                return ""
            # Replace common OCR confusions
            normalized = name.upper()
            # O and 0 are often confused - normalize to 0 in numeric contexts
            # Pattern: letter followed by O followed by digits -> replace O with 0
            normalized = re.sub(r'([A-Z])O(\d)', r'\g<1>0\2', normalized)
            # Pattern: digit followed by O -> replace O with 0
            normalized = re.sub(r'(\d)O', r'\g<1>0', normalized)
            return normalized
        
        if sample_order:
            # Create mapping from batch name to order index
            batch_order = {}
            for i, batch in enumerate(sample_order):
                # Normalize batch name (handle "LTA7TA1003/BS25080738" -> "LTA7TA1003")
                batch_key = batch.split('/')[0].upper()
                batch_order[normalize_batch_name(batch_key)] = i
            print(f"  Using sample order from Excel ({len(sample_order)} samples)")
        else:
            # Create order from unique batches found in images
            unique_batches = sorted(set(image_batches.values()))
            batch_order = {b.upper(): i for i, b in enumerate(unique_batches)}
            print(f"  Using alphabetical batch order ({len(unique_batches)} unique batches)")
        
        def get_batch_order(batch_name):
            """Get sort order for a batch name."""
            if batch_name is None:
                return 9999
            
            batch_normalized = normalize_batch_name(batch_name)
            
            # Direct match with normalized name
            if batch_normalized in batch_order:
                return batch_order[batch_normalized]
            
            # Try matching with normalized keys
            for key, order in batch_order.items():
                key_normalized = normalize_batch_name(key)
                if batch_normalized == key_normalized:
                    return order
                # Partial match
                if batch_normalized in key_normalized or key_normalized in batch_normalized:
                    return order
            
            return 9999
        
        # Step 3: Group images by section based on document structure
        # Read document.xml from the ORIGINAL template (not the modified output) to find section boundaries
        section_source = template_path if template_path else docx_path
        with zipfile.ZipFile(section_source, 'r') as z:
            doc_content_for_sections = z.read('word/document.xml').decode('utf-8')
            rels_content_for_sections = z.read('word/_rels/document.xml.rels').decode('utf-8')
        
        # Find section markers in the morphology results area
        # Look for "Images from particles" headers which mark the start of each section
        
        # Find all "Images from particles" occurrences
        img_section_matches = list(re.finditer(r'Images from particles', doc_content_for_sections, re.IGNORECASE))
        
        if len(img_section_matches) >= 2:
            # First occurrence is small particles (<100µm ECD)
            # Second occurrence is large particles (≥100µm Length)
            small_section_start = img_section_matches[0].start()
            large_section_start = img_section_matches[1].start()
        elif len(img_section_matches) == 1:
            # Only one section found - try to find the other by looking for "≥100 µm Length"
            small_section_start = img_section_matches[0].start()
            large_match = re.search(r'≥100\s*µm\s*Length', doc_content_for_sections[small_section_start:], re.IGNORECASE)
            if large_match:
                large_section_start = small_section_start + large_match.start()
            else:
                large_section_start = len(doc_content_for_sections)
        else:
            # Fallback: use image positions to determine sections
            small_section_start = 0
            large_section_start = len(doc_content_for_sections) // 2
        
        print(f"  Section boundaries: Small starts at {small_section_start}, Large starts at {large_section_start}")
        
        # Use relationships already read from docx
        rid_to_img_temp = {}
        for m in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/(image\d+\.png)"', rels_content_for_sections):
            rid_to_img_temp[m.group(1)] = m.group(2)
        
        # Find image positions in document
        img_positions = {}
        for m in re.finditer(r'r:embed="(rId\d+)"', doc_content_for_sections):
            rid = m.group(1)
            img = rid_to_img_temp.get(rid)
            if img and img in image_batches:
                img_positions[img] = m.start()
        
        # Classify images by section
        small_images = {}  # <100µm ECD
        large_images = {}  # ≥100µm Length
        
        for img_name, batch in image_batches.items():
            pos = img_positions.get(img_name, 0)
            if pos >= small_section_start and pos < large_section_start:
                small_images[img_name] = batch
            elif pos >= large_section_start:
                large_images[img_name] = batch
        
        print(f"  Found {len(small_images)} small particle images (<100µm ECD)")
        print(f"  Found {len(large_images)} large particle images (≥100µm Length)")
        
        # Step 4: Extract the docx
        extract_dir = os.path.join(temp_dir, 'docx_extracted')
        with zipfile.ZipFile(docx_path, 'r') as z:
            z.extractall(extract_dir)
        
        # Step 5: Read the relationships file to map rId to image
        rels_path = os.path.join(extract_dir, 'word', '_rels', 'document.xml.rels')
        with open(rels_path, 'r', encoding='utf-8') as f:
            rels_content = f.read()
        
        img_to_rid = {}
        rid_to_img = {}
        rid_pattern = re.compile(r'<Relationship[^>]*Id="(rId\d+)"[^>]*Target="media/(image\d+\.[a-z]+)"[^>]*/>')
        
        for match in rid_pattern.finditer(rels_content):
            rid = match.group(1)
            img_name = match.group(2)
            img_to_rid[img_name] = rid
            rid_to_img[rid] = img_name
        
        # Step 6: Read document.xml
        doc_path = os.path.join(extract_dir, 'word', 'document.xml')
        with open(doc_path, 'r', encoding='utf-8') as f:
            doc_content = f.read()
        
        # Step 7: Find all drawing elements with their rId
        drawing_pattern = re.compile(r'(<wp:(?:inline|anchor)[^>]*>.*?</wp:(?:inline|anchor)>)', re.DOTALL)
        
        drawings = []
        for match in drawing_pattern.finditer(doc_content):
            drawing_xml = match.group(1)
            start_pos = match.start()
            end_pos = match.end()
            
            rid_match = re.search(r'r:embed="(rId\d+)"', drawing_xml)
            if rid_match:
                rid = rid_match.group(1)
                img_name = rid_to_img.get(rid)
                if img_name and img_name in image_batches:
                    drawings.append({
                        'xml': drawing_xml,
                        'start': start_pos,
                        'end': end_pos,
                        'rid': rid,
                        'img': img_name,
                        'batch': image_batches[img_name]
                    })
        
        # Step 8: Group drawings by image type
        small_drawings = [d for d in drawings if d['img'] in small_images]
        large_drawings = [d for d in drawings if d['img'] in large_images]
        
        small_drawings.sort(key=lambda x: x['start'])
        large_drawings.sort(key=lambda x: x['start'])
        
        # Step 9: Create reordered versions
        def reorder_drawings(drawings_list):
            """Reorder drawings by batch order, preserving positions."""
            if len(drawings_list) <= 1:
                return []
            
            # Sort drawings by batch order
            sorted_by_batch = sorted(drawings_list, key=lambda x: get_batch_order(x['batch']))
            
            swaps = []
            for i, original in enumerate(drawings_list):
                desired = sorted_by_batch[i]
                if original['img'] != desired['img']:
                    swaps.append({
                        'position': i,
                        'original_img': original['img'],
                        'new_img': desired['img'],
                        'original_batch': original['batch'],
                        'new_batch': desired['batch'],
                        'original_start': original['start'],
                        'original_end': original['end'],
                        'new_xml': desired['xml']
                    })
            
            return swaps
        
        small_swaps = reorder_drawings(small_drawings)
        large_swaps = reorder_drawings(large_drawings)
        
        # Step 10: Apply swaps
        all_replacements = []
        
        if small_swaps:
            print("  Reordering small particle images...")
            for swap in small_swaps:
                print(f"    Position {swap['position']+1}: {swap['new_batch']} (was {swap['original_batch']})")
                all_replacements.append({
                    'start': swap['original_start'],
                    'end': swap['original_end'],
                    'new_xml': swap['new_xml']
                })
        
        if large_swaps:
            print("  Reordering large particle images...")
            for swap in large_swaps:
                print(f"    Position {swap['position']+1}: {swap['new_batch']} (was {swap['original_batch']})")
                all_replacements.append({
                    'start': swap['original_start'],
                    'end': swap['original_end'],
                    'new_xml': swap['new_xml']
                })
        
        all_replacements.sort(key=lambda x: x['start'], reverse=True)
        
        modified_doc = doc_content
        for repl in all_replacements:
            modified_doc = modified_doc[:repl['start']] + repl['new_xml'] + modified_doc[repl['end']:]
        
        # Step 11: Write modified document.xml
        with open(doc_path, 'w', encoding='utf-8') as f:
            f.write(modified_doc)
        
        # Step 12: Repackage the docx
        print("  Saving reordered document...")
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, extract_dir)
                    z.write(file_path, arc_name)
        
        print("  Image reordering complete!")
        print("  (Original image sizes and quality preserved)")
        return True
        
    except Exception as e:
        print(f"  ERROR during image reordering: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def add_rows_to_flowcam_table(table, num_rows_needed: int):
    """
    Add rows to the FlowCam table to accommodate all samples.
    
    The table has 3 header rows and 2 data rows by default.
    We need to add more data rows if there are more than 2 samples.
    """
    current_data_rows = len(table.rows) - 3  # Subtract 3 header rows
    rows_to_add = num_rows_needed - current_data_rows
    
    if rows_to_add <= 0:
        return
    
    # Get the last data row as template
    template_row = table.rows[-1]
    
    for _ in range(rows_to_add):
        # Copy the template row
        new_row_xml = copy_row(template_row)
        # Add to table
        table._tbl.append(new_row_xml)


def fill_flowcam_table(table, samples_df: pd.DataFrame):
    """
    Fill the FlowCam quantification table with sample data.
    
    Args:
        table: The FlowCam Word table
        samples_df: DataFrame with sample data from 'To Word 2'
    """
    num_samples = len(samples_df)
    
    # Add rows if needed (3 header rows + data rows)
    add_rows_to_flowcam_table(table, num_samples)
    
    # Fill each sample row (starting from row 3, after headers)
    for sample_idx in range(num_samples):
        row_idx = sample_idx + 3  # Skip 3 header rows
        sample = samples_df.iloc[sample_idx]
        
        # Set JPL Sample # (column 0)
        try:
            table.rows[row_idx].cells[0].text = str(sample_idx + 1)
        except IndexError:
            continue
        
        # Fill mapped columns
        for excel_col, table_col in FLOWCAM_COLUMN_MAP.items():
            if excel_col == "JPL Sample #":
                continue  # Already handled above
                
            value = None
            for col in samples_df.columns:
                if excel_col.lower() in col.lower() or col.lower() in excel_col.lower():
                    value = sample.get(col)
                    break
            
            if value is not None and pd.notna(value):
                try:
                    table.rows[row_idx].cells[table_col].text = str(value)
                except IndexError:
                    pass
        
        # Clear particle concentration and count columns (to be filled later)
        for col_idx in PARTICLE_CONCENTRATION_COLS + PARTICLE_COUNT_COLS:
            try:
                table.rows[row_idx].cells[col_idx].text = ""
            except IndexError:
                pass


def create_additional_table(doc: Document, template_table, samples_df: pd.DataFrame, 
                           start_sample: int, table_number: int):
    """
    Create an additional sample table for samples beyond the first 6.
    
    Returns the new table element.
    """
    # Copy the template table structure
    new_tbl = copy_table(template_table)
    
    # Create a new table from the copied XML
    new_table = doc.add_table(rows=0, cols=0)
    new_table._tbl = new_tbl
    
    # Update sample numbers in header row
    num_samples = min(6, len(samples_df) - start_sample)
    header_row = new_table.rows[0]
    
    for i in range(num_samples):
        sample_num = start_sample + i + 1
        try:
            header_row.cells[i + 1].text = str(sample_num)
        except IndexError:
            pass
    
    # Clear remaining header cells if fewer than 6 samples
    for i in range(num_samples, 6):
        try:
            header_row.cells[i + 1].text = ""
        except IndexError:
            pass
    
    # Clear all data cells first
    for row_idx in range(1, len(new_table.rows)):
        for col_idx in range(1, 7):
            try:
                new_table.rows[row_idx].cells[col_idx].text = ""
            except IndexError:
                pass
    
    # Fill with new data
    fill_sample_table(new_table, samples_df, start_sample)
    
    return new_table


def fill_report(excel_path: str, template_path: str, output_path: str, 
                particle_data_path: str = None, reorder_images: bool = False,
                footer_filename: str = None):
    """
    Main function to fill the report template with Excel data.
    
    Args:
        excel_path: Path to the Excel file
        template_path: Path to the Word template
        output_path: Path for the output Word document
        particle_data_path: Path to folder containing particle data CSVs (optional)
        reorder_images: Whether to reorder morphology images by sample sequence
        footer_filename: Filename to display in footer (e.g., "JPL25_0180_Report.docx")
    """
    # Open template
    print(f"Opening template: {template_path}")
    doc = Document(template_path)
    
    # ===== PART 1: Fill Sample Information Table from 'To Word 1' =====
    print(f"\n--- Reading 'To Word 1' data from: {excel_path} ---")
    try:
        samples_df_1 = read_excel_data_transposed(excel_path, "To Word 1")
        num_samples_1 = len(samples_df_1)
        print(f"Found {num_samples_1} samples in 'To Word 1'")
        
        print("\nSample data (To Word 1):")
        print(samples_df_1.to_string())
        print()
        
        # Get the sample table
        sample_table = get_sample_table(doc)
        if sample_table is None:
            print("WARNING: Could not find sample table in template")
        else:
            # Fill the first table (samples 1-6)
            print("Filling sample information table (samples 1-6)...")
            fill_sample_table(sample_table, samples_df_1, start_sample=0)
            
            # If more than 6 samples, create additional tables
            if num_samples_1 > 6:
                parent_cell = get_sample_table_parent_cell(doc)
                num_additional_tables = (num_samples_1 - 1) // 6
                
                for table_num in range(1, num_additional_tables + 1):
                    start_idx = table_num * 6
                    if start_idx >= num_samples_1:
                        break
                        
                    print(f"Creating additional table for samples {start_idx + 1}-{min(start_idx + 6, num_samples_1)}...")
                    
                    parent_cell.add_paragraph()
                    new_table = create_additional_table(doc, sample_table, samples_df_1, 
                                                        start_idx, table_num + 1)
                    parent_cell._element.append(new_table._tbl)
    except Exception as e:
        print(f"WARNING: Could not read 'To Word 1': {e}")
    
    # ===== PART 2: Fill FlowCam Table from 'To Word 2' =====
    print(f"\n--- Reading 'To Word 2' data from: {excel_path} ---")
    flowcam_table = None
    sample_order_list = None  # For matching particle data folders to samples
    try:
        samples_df_2 = read_excel_data_rows(excel_path, "To Word 2")
        num_samples_2 = len(samples_df_2)
        print(f"Found {num_samples_2} samples in 'To Word 2'")
        
        print("\nSample data (To Word 2):")
        print(samples_df_2.to_string())
        print()
        
        # Extract batch/lot names for particle data folder matching
        lot_col = None
        for col in samples_df_2.columns:
            if 'lot' in col.lower() or 'formulation' in col.lower():
                lot_col = col
                break
        if lot_col:
            sample_order_list = samples_df_2[lot_col].tolist()
            print(f"Sample order for folder matching: {sample_order_list[:3]}... ({len(sample_order_list)} total)")
        
        # Get the FlowCam table
        flowcam_table = get_flowcam_table(doc)
        
        # Fill the FlowCam table
        print(f"Filling FlowCam quantification table (Table 1) with {num_samples_2} samples...")
        fill_flowcam_table(flowcam_table, samples_df_2)
        print("  - Sample info columns filled")
        
    except Exception as e:
        print(f"WARNING: Could not read 'To Word 2': {e}")
    
    # ===== PART 3: Fill Particle Data from CSV files =====
    if particle_data_path and flowcam_table:
        print(f"\n--- Reading particle data from: {particle_data_path} ---")
        
        particle_data = find_particle_data_folders(particle_data_path, sample_order_list)
        
        if particle_data:
            print(f"\nFound particle data for {len(particle_data)} samples")
            print("Filling particle concentration and count columns...")
            fill_particle_data_in_table(flowcam_table, particle_data)
            print("  - Particle concentration (#/mL) columns filled: ≥3, ≥5, ≥10, ≥25, ≥50")
            print("  - Particle count (#/container) columns filled: ≥100, ≥150")
        else:
            print("WARNING: No particle data folders found")
    elif not particle_data_path:
        print("\n--- No particle data path provided ---")
        print("  Particle concentration and count columns left empty")
    
    # Save the document
    print(f"\nSaving output to: {output_path}")
    doc.save(output_path)
    
    # ===== PART 3.5: Fix Orientation dropdowns =====
    # Extract orientation values from the Excel data
    orientation_values = []
    try:
        for col in samples_df_1.columns:
            if 'orientation' in col.lower():
                orientation_values = samples_df_1[col].tolist()
                break
        
        if orientation_values:
            print(f"\n--- Fixing Orientation dropdown fields ---")
            remove_dropdown_from_orientation_row(output_path, orientation_values)
    except Exception as e:
        print(f"  WARNING: Could not fix Orientation fields: {e}")
    
    # ===== PART 3.6: Fill EP/Seidenader row and update FlowCam header =====
    # Extract EP and Seidenader values from the Excel data
    ep_values = []
    seidenader_values = []
    try:
        for col in samples_df_1.columns:
            col_lower = col.lower()
            if 'ep result' in col_lower or col_lower == 'ep result':
                ep_values = samples_df_1[col].tolist()
            elif 'seidenader' in col_lower:
                seidenader_values = samples_df_1[col].tolist()
        
        if ep_values or seidenader_values:
            print(f"\n--- Filling EP/Seidenader row ---")
            selected_option, fill_values = fill_ep_seidenader_row(output_path, ep_values, seidenader_values)
            
            # Update FlowCam table header and fill Visual Inspection values
            if selected_option:
                print(f"\n--- Updating FlowCam Visual Inspection column ---")
                update_flowcam_visual_inspection_header(output_path, selected_option, fill_values)
    except Exception as e:
        print(f"  WARNING: Could not fill EP/Seidenader row: {e}")
    
    # ===== PART 4: Reorder Morphology Images (if requested) =====
    if reorder_images:
        reorder_morphology_images(output_path, output_path, sample_order_list, template_path)
    
    # ===== PART 5: Update Footer Filename (if provided) =====
    if footer_filename:
        update_footer_filename(output_path, footer_filename)
    
    print("\nDone!")
    
    return True


def main():
    """Command-line interface."""
    reorder_images = False
    
    # Check for --reorder-images flag
    if '--reorder-images' in sys.argv:
        reorder_images = True
        sys.argv.remove('--reorder-images')
    
    if len(sys.argv) < 2:
        # Default paths for testing
        excel_path = "Gsheet Copy (5).xlsx"
        template_path = "JPLXX_XXXX_Report_V1.1.docx"
        output_path = "JPLXX_Filled_Report.docx"
        particle_data_path = "."  # Current directory
    elif len(sys.argv) == 4:
        excel_path = sys.argv[1]
        template_path = sys.argv[2]
        output_path = sys.argv[3]
        particle_data_path = None
    elif len(sys.argv) == 5:
        excel_path = sys.argv[1]
        template_path = sys.argv[2]
        output_path = sys.argv[3]
        particle_data_path = sys.argv[4]
    else:
        print("Usage: python fill_report.py <excel_file> <template_file> <output_file> [particle_data_folder] [--reorder-images]")
        print("   or: python fill_report.py  (uses default file names and current directory)")
        print("")
        print("Options:")
        print("  --reorder-images  Reorder morphology images by sample sequence (F1, F2, F3...)")
        sys.exit(1)
    
    # Validate files exist
    if not os.path.exists(excel_path):
        print(f"ERROR: Excel file not found: {excel_path}")
        sys.exit(1)
    
    if not os.path.exists(template_path):
        print(f"ERROR: Template file not found: {template_path}")
        sys.exit(1)
    
    # Prompt for footer filename
    print("\n=== JPL Report Generator ===\n")
    footer_filename = input("Enter the filename for the footer (e.g., JPL25_0172_Report_V1.docx): ").strip()
    if not footer_filename:
        # Use output filename as default
        footer_filename = os.path.basename(output_path)
        print(f"  Using output filename: {footer_filename}")
    print()
    
    success = fill_report(excel_path, template_path, output_path, particle_data_path, reorder_images, footer_filename)
    sys.exit(0 if success else 1)


def update_footer_filename(docx_path: str, new_filename: str) -> bool:
    """
    Update the filename in all footers of the document and remove highlight.
    
    The footer contains a FILENAME field with highlighted text showing the document name.
    This function updates that text to the new filename and removes the yellow highlight.
    
    Args:
        docx_path: Path to the Word document
        new_filename: New filename to display (e.g., "JPL25_0180_Report_V1.docx")
        
    Returns:
        True if successful, False otherwise
    """
    print(f"\n--- Updating footer filename to: {new_filename} ---")
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extract the docx
        extract_dir = os.path.join(temp_dir, 'docx_extracted')
        with zipfile.ZipFile(docx_path, 'r') as z:
            z.extractall(extract_dir)
        
        # Find and update all footer files
        word_dir = os.path.join(extract_dir, 'word')
        footer_files = glob.glob(os.path.join(word_dir, 'footer*.xml'))
        
        updated_count = 0
        for footer_path in footer_files:
            with open(footer_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            original_content = content
            
            # Replace any .docx filename pattern that appears to be a report name
            # Pattern: text containing JPL or Report followed by .docx
            pattern = r'(<w:t[^>]*>)([^<]*(?:JPL|Report)[^<]*\.docx)(</w:t>)'
            
            def replace_filename(match):
                return f'{match.group(1)}{new_filename}{match.group(3)}'
            
            content = re.sub(pattern, replace_filename, content, flags=re.IGNORECASE)
            
            # Remove yellow highlight from the filename area
            # Pattern: <w:highlight w:val="yellow"/>
            content = re.sub(r'<w:highlight\s+w:val="yellow"\s*/>', '', content)
            
            if content != original_content:
                with open(footer_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                updated_count += 1
                print(f"  Updated: {os.path.basename(footer_path)}")
        
        if updated_count == 0:
            print("  No footer filenames found to update")
        
        # Repackage the docx
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, extract_dir)
                    z.write(file_path, arc_name)
        
        print(f"  Footer filename updated in {updated_count} footer(s)")
        print(f"  Yellow highlight removed from footer")
        return True
        
    except Exception as e:
        print(f"  ERROR updating footer: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def reorder_images_only(input_path: str, output_path: str = None):
    """
    Standalone function to reorder morphology images in an existing document.
    
    Args:
        input_path: Path to input Word document
        output_path: Path for output (defaults to overwriting input)
    """
    if output_path is None:
        output_path = input_path
    
    print(f"Reordering images in: {input_path}")
    success = reorder_morphology_images(input_path, output_path)
    
    if success:
        print(f"Output saved to: {output_path}")
    
    return success


if __name__ == "__main__":
    # Check if running in reorder-only mode
    if len(sys.argv) >= 2 and sys.argv[1] == '--reorder-only':
        if len(sys.argv) < 3:
            print("Usage: python fill_report.py --reorder-only <input_docx> [output_docx]")
            sys.exit(1)
        
        input_path = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) > 3 else None
        
        if not os.path.exists(input_path):
            print(f"ERROR: File not found: {input_path}")
            sys.exit(1)
        
        success = reorder_images_only(input_path, output_path)
        sys.exit(0 if success else 1)
    else:
        main()
